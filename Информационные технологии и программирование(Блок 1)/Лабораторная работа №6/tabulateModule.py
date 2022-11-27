from calcModule import *
import docx
import openpyxl
from openpyxl import Workbook
doc = docx.Document()
doc.add_paragraph('Задание')
par1 = doc.add_paragraph('Выполнить расчет функции y = 4x^4+ x-10*x^2-30*x-2 5 на промежутке [1.9; 2.1] c шагом 0.01 и вывести табулированные результаты функции на этом отрезке.')
par2 = doc.add_paragraph('Кроме того, рассчитать и вывести среднее арифметическое результатов.')
doc.add_picture('test.jpg', width = docx.shared.Cm(15))
doc.save('test.docx')

try:
    workbook = openpyxl.load_workbook('tabResult.xlsx')
except Exception:
    workbook = Workbook()
sheet = workbook.active
sheets = workbook.sheetnames
sheet = workbook.active
def tabulate(a, b, h):
    sum = 0
    x = a
    n = round((b - a)/h) + 1
    for i in range(1, n+1):
        temp = calc(x)
        sum += temp
        print("x= ", '{0:.2f}'.format(x), " y = ", '{0:.3f}'.format(temp))
        cell = sheet.cell(row=i + 1, column=1)
        cell.value = x
        cell = sheet.cell(row=i + 1, column=2)
        cell.value = temp
        x += h
    avg = sum/n
    print("Среднее значение функции на промежутке ", avg)
    workbook.save('tabResult.xlsx')
